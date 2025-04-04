#!/usr/bin/env python3
import numpy as np
import matplotlib
# Use headless backend for the parameter sweep
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.animation import FuncAnimation, FFMpegWriter
from scipy.fft import fftn, ifftn, fftfreq
from scipy.interpolate import RegularGridInterpolator
from scipy.ndimage import gaussian_filter
import csv, os, datetime, imageio, itertools, sys
from docx import Document
from docx.shared import Pt

print("Current working directory:", os.getcwd())

###############################################################################
# 1. CREATE RESULTS FOLDER
###############################################################################
def create_results_folder():
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = f"run_results_{timestamp}"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    return folder_name, timestamp

results_folder, run_timestamp = create_results_folder()
print(f"[INFO] Results will be saved in folder: {results_folder}")

###############################################################################
# 2. HELPER FUNCTIONS
###############################################################################
def gaussian_3d(X, Y, Z, x0, y0, z0, sigma):
    """
    Return a normalized 3D Gaussian on (X,Y,Z) centered at (x0,y0,z0) with width sigma.
    (Used to approximate the collapse event profile.)
    """
    norm = (2 * np.pi * sigma**2) ** (3/2)
    return np.exp(-((X - x0)**2 + (Y - y0)**2 + (Z - z0)**2) / (2 * sigma**2)) / norm

def solve_poisson(rho, dx, G, relativistic_factor=0.0):
    """
    Solve the Poisson equation in 2D:
      ∇²Φ(x,y) = 4πG ρ(x,y)
    using FFT with periodic boundary conditions.
    A crude relativistic correction multiplies the potential.
    """
    ny, nx = rho.shape
    rho_k = fftn(rho)
    kx = 2 * np.pi * fftfreq(nx, d=dx)
    ky = 2 * np.pi * fftfreq(ny, d=dx)
    KX, KY = np.meshgrid(kx, ky)
    k_squared = KX**2 + KY**2
    k_squared[0, 0] = 1.0
    phi_k = -4 * np.pi * G * rho_k / k_squared
    phi_k[0, 0] = 0.0
    phi = np.real(ifftn(phi_k))
    phi *= (1 + relativistic_factor)
    return phi

def interpolate_gradient(phi, grid, dx):
    """
    Compute the gradient of φ and return a list of RegularGridInterpolator objects for each component.
    """
    grad = np.gradient(phi, dx, edge_order=2)
    interp_funcs = []
    for comp in grad:
        interp = RegularGridInterpolator(grid, comp, bounds_error=False, fill_value=None)
        interp_funcs.append(interp)
    return interp_funcs

def apply_periodic_bc(pos, L):
    """
    Apply periodic boundary conditions for a position array pos in domain [-L/2, L/2].
    """
    return ((pos + L/2) % L) - L/2

def compute_power_spectrum(field):
    """
    Compute the azimuthally averaged 2D power spectrum of a 2D field.
    Returns the radial power profile.
    """
    F = np.fft.fftshift(np.fft.fft2(field))
    psd2D = np.abs(F)**2
    ny, nx = field.shape
    y, x = np.indices((ny, nx))
    center = np.array([(ny-1)/2, (nx-1)/2])
    r = np.sqrt((x-center[1])**2 + (y-center[0])**2)
    r = r.astype(int)
    tbin = np.bincount(r.ravel(), psd2D.ravel())
    nr = np.bincount(r.ravel())
    radial_prof = tbin / (nr + 1e-8)
    return radial_prof

def estimate_noise_exponent(psd, fit_range=(1,20)):
    """
    Estimate the power-law exponent of the noise spectrum by fitting log10(power) vs log10(radial bin)
    over the specified range. Returns the slope and intercept.
    """
    bins = np.arange(len(psd))
    idx = (bins >= fit_range[0]) & (bins < fit_range[1])
    x = np.log10(bins[idx] + 1e-12)
    y = np.log10(psd[idx] + 1e-12)
    if len(x) < 2:
        return None, None
    coeffs = np.polyfit(x, y, 1)
    return coeffs  # slope, intercept

###############################################################################
# 3. FIELD SIMULATION FUNCTION (Advanced Model)
###############################################################################
def run_field_simulation(
    collapse_rate=0.5,
    collapse_sigma=0.2,
    collapse_amplitude=1.0,
    continuous_noise_amplitude=0.01,
    density_decay=0.99,
    G=1.0,
    L=10.0,
    N=256,
    steps_per_cycle=50,
    num_cycles=2,
    dt=0.05,
    relativistic_factor=0.0,
    m=1.0
):
    """
    Simulate a real scalar field φ(x,y,t) in 2D with collapse-like dynamics.
    The field obeys a modified Klein–Gordon equation with collapse and noise:
      (φ_{t+1} - 2φ_t + φ_{t-1})/dt^2 = ∇²φ_t - m^2 φ_t - collapse_rate*(φ_t - φ_avg) + sqrt(collapse_rate)*η_t,
    with periodic boundary conditions.
    After T = steps_per_cycle * num_cycles steps, the energy density is computed,
    and the gravitational potential is obtained by solving ∇²Φ = 4πG ρ.
    Returns the noise exponent (slope) of the azimuthally averaged power spectrum of Φ.
    """
    T = steps_per_cycle * num_cycles
    x = np.linspace(-L/2, L/2, N, endpoint=False)
    y = np.linspace(-L/2, L/2, N, endpoint=False)
    dx = L / N
    # Initialize field with small random perturbations.
    phi_prev = 0.01 * np.random.randn(N, N)
    phi = 0.01 * np.random.randn(N, N)
    
    # Time evolution using a leapfrog scheme.
    for t in range(T):
        phi_avg = np.mean(phi)
        noise = continuous_noise_amplitude * np.random.randn(N, N)
        noise = gaussian_filter(noise, sigma=collapse_sigma/dx)
        laplacian_phi = (np.roll(phi, 1, axis=0) + np.roll(phi, -1, axis=0) +
                         np.roll(phi, 1, axis=1) + np.roll(phi, -1, axis=1) - 4*phi) / (dx**2)
        phi_next = (2*phi - phi_prev + dt**2 * (
            laplacian_phi - m**2 * phi - collapse_rate*(phi - phi_avg) + np.sqrt(collapse_rate)*noise
        ))
        phi_prev = phi
        phi = phi_next

    # Compute energy density as effective mass density.
    phi_t = (phi - phi_prev) / dt
    grad_phi_x = (np.roll(phi, -1, axis=0) - np.roll(phi, 1, axis=0))/(2*dx)
    grad_phi_y = (np.roll(phi, -1, axis=1) - np.roll(phi, 1, axis=1))/(2*dx)
    energy_density = 0.5 * (phi_t**2 + grad_phi_x**2 + grad_phi_y**2 + m**2 * phi**2)
    
    # Solve Poisson equation for gravitational potential Φ.
    phi_grav = solve_poisson(energy_density, dx, G, relativistic_factor=relativistic_factor)
    
    # Compute power spectrum of the gravitational potential.
    psd = compute_power_spectrum(phi_grav)
    slope, intercept = estimate_noise_exponent(psd, fit_range=(1, int(0.2*N)))
    if slope is None:
        print("[WARN] Could not estimate noise exponent.")
        return None

    print(f"[INFO] Field simulation complete. Estimated noise exponent (slope) = {slope:.3f}")
    return slope

###############################################################################
# 4. PARAMETER SWEEP FUNCTION
###############################################################################
def param_sweep():
    """
    Iterates over specified parameter ranges, runs the field simulation in headless mode
    for each combination, and saves the noise exponent (slope) to a CSV file.
    Also produces a dynamic DOCX report evaluating the results.
    """
    from docx import Document
    from docx.shared import Pt

    # Define parameter ranges.
    collapse_rates = [0.1, 0.3, 0.5]
    collapse_sigmas = [0.1, 0.2]
    amplitudes = [0.5, 1.0]
    noise_amplitudes = [0.005, 0.01]
    density_decays = [0.99, 0.95]
    relativistic_factors = [0.0, 0.01]

    # Fixed simulation parameters for the sweep.
    G = 1.0
    L = 10.0
    N = 64  # lower resolution for sweep speed
    steps_per_cycle = 50
    num_cycles = 2
    dt = 0.05

    results = []
    total_combos = (len(collapse_rates) * len(collapse_sigmas) * len(amplitudes) *
                    len(noise_amplitudes) * len(density_decays) * len(relativistic_factors))
    combo_counter = 0
    print(f"[INFO] Starting parameter sweep over {total_combos} combinations...")
    sys.stdout.flush()

    for cr, cs, amp, na, dd, rf in itertools.product(collapse_rates, collapse_sigmas, amplitudes, noise_amplitudes, density_decays, relativistic_factors):
        combo_counter += 1
        print(f"[INFO] Running combo {combo_counter}/{total_combos}: cr={cr}, cs={cs}, amp={amp}, na={na}, dd={dd}, rf={rf}")
        sys.stdout.flush()
        try:
            slope_val = run_field_simulation(
                collapse_rate=cr,
                collapse_sigma=cs,
                collapse_amplitude=amp,
                continuous_noise_amplitude=na,
                density_decay=dd,
                G=G,
                L=L,
                N=N,
                steps_per_cycle=steps_per_cycle,
                num_cycles=num_cycles,
                dt=dt,
                relativistic_factor=rf,
                m=1.0
            )
        except Exception as e:
            print(f"[ERROR] Simulation failed for combo cr={cr}, cs={cs}, amp={amp}, na={na}, dd={dd}, rf={rf}: {e}")
            slope_val = None

        results.append((cr, cs, amp, na, dd, rf, slope_val))
        print(f"[INFO] Combo {combo_counter}: slope = {slope_val}")
        sys.stdout.flush()

    # Save results CSV.
    csv_filename = os.path.join(results_folder, f"param_sweep_results_{run_timestamp}.csv")
    with open(csv_filename, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["collapse_rate", "collapse_sigma", "collapse_amplitude", "noise_amplitude", "density_decay", "relativistic_factor", "slope"])
        for row in results:
            writer.writerow(row)
    print(f"[INFO] Parameter sweep results saved in {csv_filename}")

    # Dynamic evaluation of results.
    valid_slopes = [s for (_,_,_,_,_,_,s) in results if s is not None]
    if valid_slopes:
        avg_slope = np.mean(valid_slopes)
        std_slope = np.std(valid_slopes)
    else:
        avg_slope, std_slope = None, None

    # Produce DOCX report.
    doc = Document()
    doc.add_heading("Parameter Sweep Report: Emergent Gravity from Quantum Collapse", 0)
    doc.add_paragraph(f"Sweep results are saved in CSV file: {csv_filename}")
    doc.add_heading("Parameter Ranges", level=1)
    param_info = (
        f"collapse_rates = {collapse_rates}\n"
        f"collapse_sigmas = {collapse_sigmas}\n"
        f"amplitudes = {amplitudes}\n"
        f"noise_amplitudes = {noise_amplitudes}\n"
        f"density_decays = {density_decays}\n"
        f"relativistic_factors = {relativistic_factors}\n"
        f"Fixed simulation parameters: G={G}, L={L}, N={N}, steps_per_cycle={steps_per_cycle}, num_cycles={num_cycles}, dt={dt}\n"
    )
    p = doc.add_paragraph(param_info)
    p.style.font.size = Pt(11)
    doc.add_heading("Results Evaluation", level=1)
    if avg_slope is not None:
        eval_text = (
            f"The average noise exponent (slope) over the valid parameter combinations is {avg_slope:.3f} "
            f"with a standard deviation of {std_slope:.3f}. A steep negative slope (around -5) suggests that small-scale fluctuations "
            "in the gravitational potential are strongly suppressed, yielding a coherent large-scale field—consistent with the hypothesis. "
            "If many parameter combinations produce slopes near -5, this would provide encouraging evidence that quantum collapse dynamics "
            "could be responsible for an emergent gravitational field. Conversely, if the slopes are significantly less steep, the emergent effect "
            "may be weaker. Further high-resolution simulations and comparisons with experimental data are needed to confirm these trends."
        )
    else:
        eval_text = "No valid noise exponent could be estimated from the current parameter sweep. This may indicate that the simulation settings or "
        eval_text += "the model need further refinement."
    doc.add_paragraph(eval_text)
    doc.add_heading("Next Steps", level=1)
    next_steps = (
        "1) Increase simulation resolution (N) and duration (steps_per_cycle, num_cycles) for robust statistics.\n"
        "2) Identify parameter combinations that consistently yield slopes near -5 and refine those further.\n"
        "3) Run control simulations with independently generated potentials to avoid circularity.\n"
        "4) Compare the predicted noise spectrum with experimental data from short-range gravity experiments or gravitational-wave detectors.\n"
    )
    doc.add_paragraph(next_steps)
    docx_filename = os.path.join(results_folder, f"param_sweep_report_{run_timestamp}.docx")
    try:
        doc.save(docx_filename)
        print(f"[INFO] DOCX report saved as {docx_filename}")
    except Exception as e:
        print(f"[ERROR] Could not save DOCX report: {e}")

###############################################################################
# 5. OPTIONAL: VISUALIZATION SIMULATION FUNCTION
###############################################################################
def run_visualization_simulation():
    """
    Runs a simulation with full visuals (and saves a video) using a selected parameter set.
    """
    interactive_mode = True  # Set True to display visuals if possible.
    G = 1.0
    L = 10.0
    N = 256  # High resolution for visualization
    dx = L / N
    steps_per_cycle = 200
    num_cycles = 10
    total_steps = steps_per_cycle * num_cycles
    dt = 0.05

    collapse_rate = 0.5
    collapse_sigma = 0.2
    collapse_amplitude = 1.0
    continuous_noise_amplitude = 0.01
    density_decay = 0.99
    relativistic_factor = 0.01
    m = 1.0
    num_particles = 10

    x = np.linspace(-L/2, L/2, N, endpoint=False)
    grid = (x, x, x)
    X, Y, Z = np.meshgrid(x, x, x, indexing='ij')

    video_frames = []
    target_shape = None

    fig, axs = plt.subplots(2, 2, figsize=(12, 10))
    ax_rho = axs[0, 0]
    ax_phi = axs[0, 1]
    ax_time = axs[1, 0]
    ax_part = axs[1, 1]

    dens_clim = (0, collapse_amplitude * 2)
    pot_clim = (-5, 5)

    slice_index = N // 2
    dens_slice = np.zeros((N, N))
    im_rho = ax_rho.imshow(dens_slice, extent=(-L/2, L/2, -L/2, L/2), origin='lower',
                           cmap='viridis', vmin=dens_clim[0], vmax=dens_clim[1])
    ax_rho.set_title("Density Slice (z=0)")
    ax_rho.set_xlabel("x")
    ax_rho.set_ylabel("y")
    fig.colorbar(im_rho, ax=ax_rho)

    phi_slice = np.zeros((N, N))
    im_phi = ax_phi.imshow(phi_slice, extent=(-L/2, L/2, -L/2, L/2), origin='lower',
                           cmap='inferno', vmin=pot_clim[0], vmax=pot_clim[1])
    ax_phi.set_title("Potential Slice (z=0)")
    ax_phi.set_xlabel("x")
    ax_phi.set_ylabel("y")
    fig.colorbar(im_phi, ax=ax_phi)

    ax_time.set_title("Time Series: Avg. Density & Avg. Potential (per cycle)")
    ax_time.set_xlabel("Time step")
    ax_time.set_ylabel("Value")
    time_line, = ax_time.plot([], [], 'b-', label="Avg Density")
    pot_line, = ax_time.plot([], [], 'r-', label="Avg Potential")
    ax_time.legend()

    ax_part.set_title("Test Particle Positions (x-y)")
    ax_part.set_xlabel("x")
    ax_part.set_ylabel("y")
    colors = plt.cm.rainbow(np.linspace(0, 1, num_particles))
    particle_positions = np.random.uniform(-L/2, L/2, (num_particles, 3))
    particle_scatter = ax_part.scatter(particle_positions[:, 0], particle_positions[:, 1], c=colors)
    ax_part.set_xlim(-L/2, L/2)
    ax_part.set_ylim(-L/2, L/2)

    rho = np.zeros((N, N, N))
    particle_velocities = np.zeros((num_particles, 3))
    cycle_time = []
    cycle_avg_rho = []
    cycle_avg_phi = []
    global_step = 0
    current_cycle = 0

    def rk2_update(p, v, dt, acc_func):
        a1 = acc_func(p)
        p_temp = p + v * dt/2 + 0.5 * a1 * (dt/2)**2
        v_temp = v + a1 * dt/2
        a2 = acc_func(p_temp)
        v_new = v + dt * 0.5 * (a1 + a2)
        p_new = p + dt * v_new
        return p_new, v_new

    def update(frame):
        nonlocal global_step, current_cycle, rho, particle_positions, particle_velocities, target_shape

        if global_step > 0 and global_step % steps_per_cycle == 0:
            current_cycle += 1
            cycle_time.clear()
            cycle_avg_rho.clear()
            cycle_avg_phi.clear()
            rho[:] = 0.0
            particle_positions[:] = np.random.uniform(-L/2, L/2, (num_particles, 3))
            particle_velocities[:] = 0

        num_events = np.random.poisson(lam=collapse_rate)
        for _ in range(num_events):
            event_pos = np.random.uniform(-L/2, L/2, size=3)
            rho += collapse_amplitude * gaussian_3d(X, Y, Z,
                                                    event_pos[0],
                                                    event_pos[1],
                                                    event_pos[2],
                                                    collapse_sigma)
        noise = continuous_noise_amplitude * np.random.randn(N, N, N)
        noise = gaussian_filter(noise, sigma=collapse_sigma/(L/N))
        rho += noise
        rho *= density_decay

        phi = solve_poisson(rho, dx=(L/N), G=G, relativistic_factor=relativistic_factor)
        interp_grad = interpolate_gradient(phi, grid, dx=(L/N))

        def acc_func(p):
            return -np.array([interp_grad[d](p.reshape(1, 3))[0] for d in range(3)])
        for i in range(num_particles):
            pos = particle_positions[i]
            vel = particle_velocities[i]
            new_pos, new_vel = rk2_update(pos, vel, dt, acc_func)
            particle_positions[i] = apply_periodic_bc(new_pos, L)
            particle_velocities[i] = new_vel

        dens_slice = rho[:, :, N//2]
        phi_slice = phi[:, :, N//2]
        avg_rho = np.mean(dens_slice)
        avg_phi = np.mean(phi_slice)
        cycle_time.append(global_step % steps_per_cycle)
        cycle_avg_rho.append(avg_rho)
        cycle_avg_phi.append(avg_phi)

        im_rho.set_data(dens_slice)
        im_phi.set_data(phi_slice)
        time_line.set_data(np.array(cycle_time), np.array(cycle_avg_rho))
        pot_line.set_data(np.array(cycle_time), np.array(cycle_avg_phi))
        ax_time.relim()
        ax_time.autoscale_view()
        particle_scatter.set_offsets(particle_positions[:, :2])
        fig.suptitle(f"Cycle {current_cycle+1}/{num_cycles} - Step {(global_step % steps_per_cycle)+1}/{steps_per_cycle}")
        global_step += 1

        fig.canvas.draw()
        frame_image = np.frombuffer(fig.canvas.tostring_rgb(), dtype='uint8')
        frame_image = frame_image.reshape(fig.canvas.get_width_height()[::-1] + (3,))
        if target_shape is None:
            target_shape = frame_image.shape
        else:
            if frame_image.shape != target_shape:
                frame_image = frame_image[:target_shape[0], :target_shape[1], :]
        video_frames.append(frame_image)
        return im_rho, im_phi, time_line, pot_line, particle_scatter

    anim = FuncAnimation(fig, update, frames=total_steps, interval=50, blit=False, repeat=False)
    plt.tight_layout()
    if interactive_mode:
        plt.show()
    else:
        plt.close(fig)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    video_filename = os.path.join(results_folder, f"visualization_simulation_{timestamp}.mp4")
    try:
        imageio.mimsave(video_filename, video_frames, fps=20)
        print(f"[INFO] Visualization video saved as {video_filename}")
    except Exception as e:
        print(f"[ERROR] Could not save visualization video: {e}")

###############################################################################
# 6. MAIN FUNCTION
###############################################################################
def main():
    print("[INFO] Starting parameter sweep...")
    param_sweep()
    print("[INFO] Parameter sweep completed.")
    choice = input("Do you want to run the visualization simulation? (y/n): ").strip().lower()
    if choice == 'y':
        print("[INFO] Starting visualization simulation...")
        run_visualization_simulation()
        print("[INFO] Visualization simulation completed.")
    else:
        print("[INFO] Visualization simulation skipped.")

if __name__ == '__main__':
    main()

