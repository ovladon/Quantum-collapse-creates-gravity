#!/usr/bin/env python3
import numpy as np
import matplotlib
# Use headless backend for optimization
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from scipy.fft import fftn, ifftn, fftfreq
from scipy.interpolate import RegularGridInterpolator
from scipy.ndimage import gaussian_filter
import csv, os, datetime, itertools, sys
from docx import Document
from docx.shared import Pt

print("Current working directory:", os.getcwd())

###############################################################################
# 1. CREATE RESULTS FOLDER
###############################################################################
def create_results_folder():
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = f"opt_results_{timestamp}"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    return folder_name, timestamp

results_folder, run_timestamp = create_results_folder()
print(f"[INFO] Results will be saved in folder: {results_folder}")

###############################################################################
# 2. HELPER FUNCTIONS
###############################################################################
def gaussian_3d(X, Y, Z, x0, y0, z0, sigma):
    """
    Return a normalized 3D Gaussian on (X, Y, Z) centered at (x0,y0,z0) with width sigma.
    """
    norm = (2 * np.pi * sigma**2) ** (3/2)
    return np.exp(-((X - x0)**2 + (Y - y0)**2 + (Z - z0)**2) / (2 * sigma**2)) / norm

def solve_poisson(rho, dx, G, relativistic_factor=0.0):
    """
    Solve the Poisson equation in 2D:
       ∇²Φ(x,y) = 4πG ρ(x,y)
    using FFT with periodic boundary conditions.
    Apply a crude relativistic correction by multiplying the result by (1 + relativistic_factor).
    """
    ny, nx = rho.shape
    rho_k = fftn(rho)
    kx = 2 * np.pi * fftfreq(nx, d=dx)
    ky = 2 * np.pi * fftfreq(ny, d=dx)
    KX, KY = np.meshgrid(kx, ky)
    k_squared = KX**2 + KY**2
    k_squared[0, 0] = 1.0  # avoid division by zero
    phi_k = -4 * np.pi * G * rho_k / k_squared
    phi_k[0, 0] = 0.0
    phi = np.real(ifftn(phi_k))
    phi *= (1 + relativistic_factor)
    return phi

def interpolate_gradient(phi, grid, dx):
    """
    Compute the gradient of φ and return a list of RegularGridInterpolator objects for each component.
    """
    grad = np.gradient(phi, dx, edge_order=2)
    interp_funcs = []
    for comp in grad:
        interp = RegularGridInterpolator(grid, comp, bounds_error=False, fill_value=None)
        interp_funcs.append(interp)
    return interp_funcs

def apply_periodic_bc(pos, L):
    """
    Apply periodic boundary conditions for a position array pos in domain [-L/2, L/2].
    """
    return ((pos + L/2) % L) - L/2

def compute_power_spectrum(field):
    """
    Compute the azimuthally averaged 2D power spectrum of a 2D field.
    Returns the radial power profile.
    """
    F = np.fft.fftshift(np.fft.fft2(field))
    psd2D = np.abs(F)**2
    ny, nx = field.shape
    y, x = np.indices((ny, nx))
    center = np.array([(ny-1)/2, (nx-1)/2])
    r = np.sqrt((x-center[1])**2 + (y-center[0])**2)
    r = r.astype(int)
    tbin = np.bincount(r.ravel(), psd2D.ravel())
    nr = np.bincount(r.ravel())
    radial_prof = tbin / (nr + 1e-8)
    return radial_prof

def estimate_noise_exponent(psd, fit_range=(1,20)):
    """
    Estimate the power-law exponent of the noise spectrum by fitting log10(power) vs. log10(radial bin)
    over the specified range. Returns the slope and intercept.
    """
    bins = np.arange(len(psd))
    idx = (bins >= fit_range[0]) & (bins < fit_range[1])
    x = np.log10(bins[idx] + 1e-12)
    y = np.log10(psd[idx] + 1e-12)
    if len(x) < 2:
        return None, None
    coeffs = np.polyfit(x, y, 1)
    return coeffs  # slope, intercept

###############################################################################
# 3. FIELD SIMULATION FUNCTION (Advanced 2D Model)
###############################################################################
def run_field_simulation(
    collapse_rate=0.5,
    collapse_sigma=0.2,
    collapse_amplitude=1.0,
    continuous_noise_amplitude=0.01,
    density_decay=0.99,
    G=1.0,
    L=10.0,
    N=256,
    steps_per_cycle=50,
    num_cycles=2,
    dt=0.05,
    relativistic_factor=0.0,
    m=1.0
):
    """
    Simulate a real scalar field φ(x,y,t) in 2D with collapse-like dynamics.
    The field obeys a modified Klein–Gordon equation with collapse and noise:
      (φ_{t+1} - 2φ_t + φ_{t-1})/dt^2 = ∇²φ_t - m^2 φ_t - collapse_rate*(φ_t - φ_avg) + sqrt(collapse_rate)*η_t,
    with periodic boundary conditions.
    After T = steps_per_cycle * num_cycles steps, compute the energy density and solve the Poisson
    equation for the gravitational potential Φ. Then, compute the azimuthally averaged power spectrum of Φ
    and return the estimated noise exponent (slope).
    """
    T = steps_per_cycle * num_cycles
    x = np.linspace(-L/2, L/2, N, endpoint=False)
    y = np.linspace(-L/2, L/2, N, endpoint=False)
    dx = L / N
    # Initialize field with small random perturbations.
    phi_prev = 0.01 * np.random.randn(N, N)
    phi = 0.01 * np.random.randn(N, N)
    
    for t in range(T):
        phi_avg = np.mean(phi)
        noise = continuous_noise_amplitude * np.random.randn(N, N)
        noise = gaussian_filter(noise, sigma=collapse_sigma/dx)
        laplacian_phi = (np.roll(phi, 1, axis=0) + np.roll(phi, -1, axis=0) +
                         np.roll(phi, 1, axis=1) + np.roll(phi, -1, axis=1) - 4*phi) / (dx**2)
        phi_next = (2*phi - phi_prev + dt**2 * (
            laplacian_phi - m**2 * phi - collapse_rate*(phi - phi_avg) + np.sqrt(collapse_rate)*noise
        ))
        phi_prev = phi
        phi = phi_next

    phi_t = (phi - phi_prev) / dt
    grad_phi_x = (np.roll(phi, -1, axis=0) - np.roll(phi, 1, axis=0))/(2*dx)
    grad_phi_y = (np.roll(phi, -1, axis=1) - np.roll(phi, 1, axis=1))/(2*dx)
    energy_density = 0.5 * (phi_t**2 + grad_phi_x**2 + grad_phi_y**2 + m**2 * phi**2)
    
    phi_grav = solve_poisson(energy_density, dx, G, relativistic_factor=relativistic_factor)
    
    psd = compute_power_spectrum(phi_grav)
    slope, intercept = estimate_noise_exponent(psd, fit_range=(1, int(0.2*N)))
    if slope is None:
        print("[WARN] Could not estimate noise exponent.")
        return None
    print(f"[INFO] Field simulation complete. Estimated noise exponent (slope) = {slope:.3f}")
    return slope

###############################################################################
# 4. EVOLUTIONARY PARAMETER OPTIMIZATION FUNCTION
###############################################################################
def optimize_parameters(num_iterations=10, samples_per_iteration=20):
    """
    Evolutionary optimization tool that runs the field simulation repeatedly.
    It samples parameter combinations from current ranges, evaluates fitness based on
    how close the noise exponent is to -5, refines the ranges iteratively, and saves CSV
    files for each iteration. Returns the best configuration found.
    """
    from docx import Document
    from docx.shared import Pt

    # Define initial parameter ranges as (min, max).
    param_ranges = {
        "collapse_rate": (0.1, 0.5),
        "collapse_sigma": (0.1, 0.2),
        "collapse_amplitude": (0.5, 1.0),
        "continuous_noise_amplitude": (0.005, 0.01),
        "density_decay": (0.95, 0.99),
        "relativistic_factor": (0.0, 0.01)
    }
    
    # Fixed simulation parameters.
    G = 1.0
    L = 10.0
    N = 64   # Lower resolution for optimization speed.
    steps_per_cycle = 50
    num_cycles = 2
    dt = 0.05
    m = 1.0

    best_configurations = []
    all_iteration_results = []
    
    def sample_parameters(ranges):
        return {k: np.random.uniform(v[0], v[1]) for k, v in ranges.items()}
    
    def fitness(slope):
        # We want slope as close to -5 as possible.
        if slope is None:
            return -np.inf
        return -abs(slope + 5)
    
    for iteration in range(num_iterations):
        iteration_results = []
        print(f"[ITERATION {iteration+1}] Current parameter ranges: {param_ranges}")
        sys.stdout.flush()
        for i in range(samples_per_iteration):
            config = sample_parameters(param_ranges)
            try:
                slope_val = run_field_simulation(
                    collapse_rate=config["collapse_rate"],
                    collapse_sigma=config["collapse_sigma"],
                    collapse_amplitude=config["collapse_amplitude"],
                    continuous_noise_amplitude=config["continuous_noise_amplitude"],
                    density_decay=config["density_decay"],
                    G=G,
                    L=L,
                    N=N,
                    steps_per_cycle=steps_per_cycle,
                    num_cycles=num_cycles,
                    dt=dt,
                    relativistic_factor=config["relativistic_factor"],
                    m=m
                )
            except Exception as e:
                print(f"[ERROR] Simulation failed for config {config}: {e}")
                slope_val = None
            fit = fitness(slope_val)
            iteration_results.append((config, fit, slope_val))
            print(f"[ITERATION {iteration+1}] Sample {i+1}: {config} -> slope: {slope_val}, fitness: {fit}")
            sys.stdout.flush()
        # Save CSV for this iteration.
        iter_csv = os.path.join(results_folder, f"iter_{iteration+1}_results_{run_timestamp}.csv")
        with open(iter_csv, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["collapse_rate", "collapse_sigma", "collapse_amplitude", "continuous_noise_amplitude",
                             "density_decay", "relativistic_factor", "slope"])
            for row in iteration_results:
                config, fit, slope_val = row
                writer.writerow([config["collapse_rate"], config["collapse_sigma"], config["collapse_amplitude"],
                                 config["continuous_noise_amplitude"], config["density_decay"],
                                 config["relativistic_factor"], slope_val])
        print(f"[INFO] Iteration {iteration+1} results saved in {iter_csv}")
        sys.stdout.flush()
        
        # Select top 20% of configurations.
        iteration_results = sorted(iteration_results, key=lambda x: x[1], reverse=True)
        top_n = max(1, int(0.2 * len(iteration_results)))
        best_configs = iteration_results[:top_n]
        best_configurations.extend(best_configs)
        all_iteration_results.append(iteration_results)
        
        # Refine parameter ranges: update each range around the average of the top configurations.
        for key in param_ranges.keys():
            vals = [cfg[0][key] for cfg in best_configs if cfg[0][key] is not None]
            if vals:
                mean_val = np.mean(vals)
                current_min, current_max = param_ranges[key]
                width = current_max - current_min
                # Narrow range to +/- 50% of the current width around the mean.
                new_min = max(current_min, mean_val - 0.5 * width / 2)
                new_max = min(current_max, mean_val + 0.5 * width / 2)
                if new_max - new_min < 1e-4:  # too narrow, revert to original
                    new_min, new_max = current_min, current_max
                param_ranges[key] = (new_min, new_max)
        print(f"[ITERATION {iteration+1}] Updated parameter ranges: {param_ranges}")
        sys.stdout.flush()
    
    # From all best configurations, select the one with highest fitness.
    best_overall = max(best_configurations, key=lambda x: x[1])
    print(f"[INFO] Best overall configuration: {best_overall[0]} with slope {best_overall[2]:.3f}")
    
    # Save final best configuration in a CSV.
    final_csv = os.path.join(results_folder, f"final_results_{run_timestamp}.csv")
    with open(final_csv, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["collapse_rate", "collapse_sigma", "collapse_amplitude", "continuous_noise_amplitude",
                         "density_decay", "relativistic_factor", "slope"])
        writer.writerow([best_overall[0]["collapse_rate"], best_overall[0]["collapse_sigma"],
                         best_overall[0]["collapse_amplitude"], best_overall[0]["continuous_noise_amplitude"],
                         best_overall[0]["density_decay"], best_overall[0]["relativistic_factor"], best_overall[2]])
    print(f"[INFO] Final optimized results saved in {final_csv}")
    sys.stdout.flush()
    return best_overall, all_iteration_results

###############################################################################
# 5. DOCX REPORT FOR OPTIMIZATION RESULTS
###############################################################################
def write_optimization_report(best_config, iteration_results):
    """
    Writes a DOCX report summarizing the optimization process and the best parameters found.
    """
    doc = Document()
    doc.add_heading("Parameter Optimization Report: Emergent Gravity from Quantum Collapse", 0)
    doc.add_heading("Final Optimized Configuration", level=1)
    best_text = "Best parameters found:\n"
    for k, v in best_config[0].items():
        best_text += f"  {k}: {v:.4f}\n"
    best_text += f"Estimated noise exponent (slope): {best_config[2]:.3f}\n"
    best_text += f"Fitness: {-abs(best_config[2] + 5):.4f}"
    doc.add_paragraph(best_text)
    
    doc.add_heading("Optimization Process Summary", level=1)
    summary = (
        "The optimization algorithm ran for multiple iterations, sampling parameter combinations and evaluating the fitness "
        "(defined as -|slope + 5|, with a target slope of -5). In each iteration, the top 20% configurations were selected to refine "
        "the parameter ranges. The following table summarizes the final iteration's best configurations:\n"
    )
    for i, res in enumerate(iteration_results[-1]):
        config, fit, slope = res
        summary += f"Config {i+1}: {config} -> slope: {slope}\n"
    doc.add_paragraph(summary)
    
    doc.add_heading("Next Steps", level=1)
    next_steps = (
        "1) Run higher-resolution simulations using the optimized parameters.\n"
        "2) Validate the robustness of the noise exponent with longer simulation durations.\n"
        "3) Run control simulations with independently generated potentials to verify the emergent noise signature.\n"
        "4) Compare the predicted noise spectrum with experimental data from precision gravity experiments.\n"
    )
    doc.add_paragraph(next_steps)
    
    report_filename = os.path.join(results_folder, f"optimization_report_{run_timestamp}.docx")
    try:
        doc.save(report_filename)
        print(f"[INFO] Optimization DOCX report saved as {report_filename}")
    except Exception as e:
        print(f"[ERROR] Could not save optimization DOCX report: {e}")

###############################################################################
# 6. MAIN FUNCTION
###############################################################################
def main():
    print("[INFO] Starting parameter optimization...")
    best_config, iteration_results = optimize_parameters(num_iterations=10, samples_per_iteration=20)
    print("[INFO] Parameter optimization completed.")
    write_optimization_report(best_config, iteration_results)
    print("[INFO] All optimization results have been saved.")
    print("Final optimized configuration:")
    print(best_config[0])
    print(f"Estimated noise exponent (slope): {best_config[2]:.3f}")

if __name__ == '__main__':
    main()

