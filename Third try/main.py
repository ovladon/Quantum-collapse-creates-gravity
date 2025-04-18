#!/usr/bin/env python3
import numpy as np
import matplotlib
# Use headless backend for parameter sweep; visuals will be optional later.
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.animation import FuncAnimation
from scipy.fft import fftn, ifftn, fftfreq
from scipy.interpolate import RegularGridInterpolator
from scipy.ndimage import gaussian_filter
import csv, os, datetime, imageio, itertools, sys
from docx import Document
from docx.shared import Pt

print("Current working directory:", os.getcwd())

###############################################################################
# 1. HELPER FUNCTIONS
###############################################################################

def gaussian_3d(X, Y, Z, x0, y0, z0, sigma):
    """
    Return a normalized 3D Gaussian on (X, Y, Z) centered at (x0, y0, z0)
    with width sigma.
    """
    norm = (2 * np.pi * sigma**2) ** (3/2)
    return np.exp(-((X - x0)**2 + (Y - y0)**2 + (Z - z0)**2) / (2 * sigma**2)) / norm

def solve_poisson(rho, dx, G):
    """
    Solve the Poisson equation ∇²Φ = 4πG ρ on a periodic grid using FFT.
    """
    N = rho.shape[0]
    rho_k = fftn(rho)
    k = 2 * np.pi * fftfreq(N, d=dx)
    KX, KY, KZ = np.meshgrid(k, k, k, indexing='ij')
    k_squared = KX**2 + KY**2 + KZ**2
    k_squared[0, 0, 0] = 1.0  # avoid division by zero
    phi_k = -4 * np.pi * G * rho_k / k_squared
    phi_k[0, 0, 0] = 0.0  # enforce zero mean
    phi = np.real(ifftn(phi_k))
    return phi

def interpolate_gradient(phi, grid, dx):
    """
    Compute the gradient of φ and return a list of RegularGridInterpolator objects
    for each component.
    """
    grad = np.gradient(phi, dx, edge_order=2)
    interp_funcs = []
    for comp in grad:
        interp = RegularGridInterpolator(grid, comp, bounds_error=False, fill_value=None)
        interp_funcs.append(interp)
    return interp_funcs

def apply_periodic_bc(pos, L):
    """
    Apply periodic boundary conditions for a position array pos in domain [-L/2, L/2].
    """
    return ((pos + L/2) % L) - L/2

def compute_power_spectrum(field):
    """
    Compute the 2D power spectrum (azimuthally averaged) of a 2D field.
    Returns the radial power profile.
    """
    F = np.fft.fftshift(np.fft.fft2(field))
    psd2D = np.abs(F)**2
    ny, nx = field.shape
    y, x = np.indices((ny, nx))
    center = np.array([(ny-1)/2, (nx-1)/2])
    r = np.sqrt((x-center[1])**2 + (y-center[0])**2)
    r = r.astype(int)  # use built-in int
    tbin = np.bincount(r.ravel(), psd2D.ravel())
    nr = np.bincount(r.ravel())
    radial_prof = tbin / (nr + 1e-8)
    return radial_prof

def estimate_noise_exponent(psd, fit_range=(1,20)):
    """
    Estimate the power-law exponent of the noise spectrum by fitting log10(power)
    vs. log10(radial bin) over the specified range. Returns the slope and intercept.
    """
    bins = np.arange(len(psd))
    idx = (bins >= fit_range[0]) & (bins < fit_range[1])
    x = np.log10(bins[idx] + 1e-12)
    y = np.log10(psd[idx] + 1e-12)
    if len(x) < 2:
        return None, None
    coeffs = np.polyfit(x, y, 1)
    return coeffs  # slope, intercept

###############################################################################
# 2. SIMULATION FUNCTION: Run a simulation and return the noise spectrum slope.
###############################################################################

def run_simulation_and_return_slope(
    collapse_rate=0.5,
    collapse_sigma=0.2,
    collapse_amplitude=1.0,
    continuous_noise_amplitude=0.01,
    density_decay=0.99,
    G=1.0,
    L=10.0,
    N=64,
    steps_per_cycle=50,
    num_cycles=2,
    dt=0.05
):
    """
    Runs a headless simulation with the given parameters.
    Returns the slope of the mid-plane potential's noise spectrum averaged over cycles.
    In case of error, returns None.
    """
    print(f"[INFO] Running simulation with parameters: collapse_rate={collapse_rate}, collapse_sigma={collapse_sigma}, "
          f"amplitude={collapse_amplitude}, noise_amplitude={continuous_noise_amplitude}, density_decay={density_decay}, "
          f"G={G}, L={L}, N={N}, steps_per_cycle={steps_per_cycle}, num_cycles={num_cycles}, dt={dt}")
    sys.stdout.flush()

    total_steps = steps_per_cycle * num_cycles
    # Create the 3D grid
    x = np.linspace(-L/2, L/2, N, endpoint=False)
    grid = (x, x, x)
    # Initialize state
    rho = np.zeros((N, N, N))
    # We'll ignore particles if only noise spectrum is of interest.
    num_particles = 10
    particle_positions = np.random.uniform(-L/2, L/2, (num_particles, 3))
    particle_velocities = np.zeros((num_particles, 3))
    cycle_psd = []  # list to hold PSD for each step in a cycle
    current_cycle = 0

    # Use a simple RK2 integrator for test particles.
    def rk2_update(p, v, dt, acc_func):
        a1 = acc_func(p)
        p_temp = p + v * dt/2 + 0.5 * a1 * (dt/2)**2
        v_temp = v + a1 * dt/2
        a2 = acc_func(p_temp)
        v_new = v + dt * 0.5 * (a1 + a2)
        p_new = p + dt * v_new
        return p_new, v_new

    for step in range(total_steps):
        # At end of cycle, reset and store PSD
        if step > 0 and step % steps_per_cycle == 0:
            current_cycle += 1
            # Average PSD for this cycle:
            avg_psd = np.mean(np.array(cycle_psd), axis=0)
            # For this demonstration, we take the PSD from the final cycle.
            final_psd = avg_psd
            cycle_psd = []
            rho[:] = 0.0
            particle_positions[:] = np.random.uniform(-L/2, L/2, (num_particles, 3))
            particle_velocities[:] = 0

        # (a) Add discrete collapse events.
        num_events = np.random.poisson(lam=collapse_rate)
        # Build a 3D grid for the simulation:
        X, Y, Z = np.meshgrid(x, x, x, indexing='ij')
        for _ in range(num_events):
            event_pos = np.random.uniform(-L/2, L/2, size=3)
            rho += collapse_amplitude * gaussian_3d(X, Y, Z,
                                                    event_pos[0],
                                                    event_pos[1],
                                                    event_pos[2],
                                                    collapse_sigma)
        # (b) Add continuous noise.
        noise = continuous_noise_amplitude * np.random.randn(N, N, N)
        noise = gaussian_filter(noise, sigma=collapse_sigma/(L/N))
        rho += noise
        # (c) Apply density decay.
        rho *= density_decay
        # (d) Solve Poisson to get potential.
        phi = solve_poisson(rho, dx=(L/N), G=G)
        # (e) Compute gradient.
        interp_grad = interpolate_gradient(phi, grid, dx=(L/N))
        # (f) (Optionally update particles, but not necessary for PSD.)
        # (g) Compute PSD of the mid-plane (z = N//2) potential.
        mid_plane = phi[:, :, N//2]
        psd_slice = compute_power_spectrum(mid_plane)
        cycle_psd.append(psd_slice)

    # After simulation, average PSD over last cycle:
    if len(cycle_psd) > 0:
        overall_psd = np.mean(np.array(cycle_psd), axis=0)
    else:
        print("[WARN] No PSD data collected.")
        return None

    slope, intercept = estimate_noise_exponent(overall_psd, fit_range=(1,20))
    if slope is None:
        print("[WARN] Noise exponent could not be estimated.")
        return None

    print(f"[INFO] Simulation complete. Estimated noise exponent (slope) = {slope:.3f}")
    return slope

###############################################################################
# 3. PARAMETER SWEEP
###############################################################################

def param_sweep():
    """
    Performs a parameter sweep over different collapse parameters.
    Writes the results (parameter combination and computed noise spectrum slope)
    to a CSV file, and produces a DOCX summary report.
    """
    from docx import Document
    from docx.shared import Pt

    # Define parameter ranges.
    collapse_rates = [0.1, 0.3, 0.5]
    collapse_sigmas = [0.1, 0.2]
    amplitudes = [0.5, 1.0]
    noise_amplitudes = [0.005, 0.01]
    density_decays = [0.99, 0.95]

    # Fixed parameters for sweep simulation.
    G = 1.0
    L = 10.0
    N = 64   # Use lower resolution for sweep speed.
    steps_per_cycle = 50
    num_cycles = 2
    dt = 0.05

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    csv_filename = f"param_sweep_results_{timestamp}.csv"
    total_combos = len(collapse_rates) * len(collapse_sigmas) * len(amplitudes) * len(noise_amplitudes) * len(density_decays)
    print(f"[INFO] Sweeping over {total_combos} parameter combinations...")

    results = []
    combo_counter = 0

    for cr, cs, amp, na, dd in itertools.product(collapse_rates, collapse_sigmas, amplitudes, noise_amplitudes, density_decays):
        combo_counter += 1
        print(f"[INFO] Running combo {combo_counter}/{total_combos}: cr={cr}, cs={cs}, amp={amp}, na={na}, dd={dd}")
        sys.stdout.flush()
        try:
            slope_val = run_simulation_and_return_slope(
                collapse_rate=cr,
                collapse_sigma=cs,
                collapse_amplitude=amp,
                continuous_noise_amplitude=na,
                density_decay=dd,
                G=G,
                L=L,
                N=N,
                steps_per_cycle=steps_per_cycle,
                num_cycles=num_cycles,
                dt=dt
            )
        except Exception as e:
            print(f"[ERROR] Simulation failed for combo cr={cr}, cs={cs}, amp={amp}, na={na}, dd={dd}: {e}")
            slope_val = None

        results.append((cr, cs, amp, na, dd, slope_val))
        print(f"[INFO] Combo {combo_counter}: slope = {slope_val}")
        sys.stdout.flush()

    # Write results to CSV.
    with open(csv_filename, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["collapse_rate", "collapse_sigma", "collapse_amplitude", "noise_amplitude", "density_decay", "slope"])
        for row in results:
            writer.writerow(row)
    print(f"[INFO] Parameter sweep results saved in {csv_filename}")

    # Produce DOCX report.
    doc = Document()
    doc.add_heading("Parameter Sweep Report: Emergent Gravity from Quantum Collapse", 0)
    doc.add_paragraph(f"Sweep results are saved in CSV file: {csv_filename}")
    doc.add_heading("Parameter Ranges", level=1)
    param_info = (
        f"collapse_rates = {collapse_rates}\n"
        f"collapse_sigmas = {collapse_sigmas}\n"
        f"amplitudes = {amplitudes}\n"
        f"noise_amplitudes = {noise_amplitudes}\n"
        f"density_decays = {density_decays}\n"
        f"Fixed sim parameters: G={G}, L={L}, N={N}, steps_per_cycle={steps_per_cycle}, num_cycles={num_cycles}, dt={dt}\n"
    )
    p = doc.add_paragraph(param_info)
    p.style.font.size = Pt(11)
    doc.add_heading("Results and Discussion", level=1)
    discussion = (
        "The parameter sweep systematically explores different combinations of collapse parameters. "
        "The noise spectrum slope (power-law exponent) is computed for each combination. A steeper negative slope "
        "(e.g., near -5) suggests that high-frequency fluctuations are strongly suppressed, which might indicate that "
        "the emergent gravitational potential is dominated by large-scale, coherent structures. Conversely, a shallower slope "
        "(e.g., near -1 or -2) would imply more small-scale noise. These results provide an independent quantitative signature "
        "that can eventually be compared with experimental data from short-range gravity tests or gravitational-wave detectors.\n"
    )
    doc.add_paragraph(discussion)
    doc.add_heading("Next Steps", level=1)
    next_steps = (
        "1) Identify parameter combinations that consistently yield slopes near -5.\n"
        "2) Increase simulation resolution and number of cycles for robust statistics.\n"
        "3) Run control simulations with independently generated potentials to guard against circularity.\n"
        "4) Compare the predicted noise spectrum with experimental data from precision gravity experiments.\n"
    )
    doc.add_paragraph(next_steps)
    try:
        docx_filename = f"param_sweep_report_{timestamp}.docx"
        doc.save(docx_filename)
        print(f"[INFO] DOCX report saved as {docx_filename}")
    except Exception as e:
        print(f"[ERROR] Could not save DOCX report: {e}")

###############################################################################
# 4. OPTIONAL: VISUALIZATION FUNCTION
###############################################################################

def run_visualization_simulation():
    """
    Runs a simulation with full visuals (and compiles a video) using a selected parameter combination.
    Change parameters below as desired.
    """
    # Set interactive mode to True so the GUI window shows.
    interactive_mode = True

    # Use parameters of interest (choose one from your sweep results that looks promising)
    G = 1.0
    L = 10.0
    N = 128
    dx = L / N
    steps_per_cycle = 200
    num_cycles = 10
    total_steps = steps_per_cycle * num_cycles
    dt = 0.05

    collapse_rate = 0.5
    collapse_sigma = 0.2
    collapse_amplitude = 1.0
    continuous_noise_amplitude = 0.01
    density_decay = 0.99

    num_particles = 10

    # Create grid.
    x = np.linspace(-L/2, L/2, N, endpoint=False)
    grid = (x, x, x)
    X, Y, Z = np.meshgrid(x, x, x, indexing='ij')

    # Containers for video frames.
    video_frames = []
    target_shape = None

    # Set up figure with 4 subplots.
    fig, axs = plt.subplots(2, 2, figsize=(12, 10))
    ax_rho = axs[0, 0]
    ax_phi = axs[0, 1]
    ax_time = axs[1, 0]
    ax_part = axs[1, 1]

    # Fixed color limits.
    dens_clim = (0, collapse_amplitude * 2)
    pot_clim = (-5, 5)

    # Prepare initial images (mid-plane slice: z=0).
    slice_index = N // 2
    dens_slice = np.zeros((N, N))
    im_rho = ax_rho.imshow(dens_slice, extent=(-L/2, L/2, -L/2, L/2), origin='lower',
                           cmap='viridis', vmin=dens_clim[0], vmax=dens_clim[1])
    ax_rho.set_title("Density Slice (z=0)")
    ax_rho.set_xlabel("x")
    ax_rho.set_ylabel("y")
    fig.colorbar(im_rho, ax=ax_rho)

    phi_slice = np.zeros((N, N))
    im_phi = ax_phi.imshow(phi_slice, extent=(-L/2, L/2, -L/2, L/2), origin='lower',
                           cmap='inferno', vmin=pot_clim[0], vmax=pot_clim[1])
    ax_phi.set_title("Potential Slice (z=0)")
    ax_phi.set_xlabel("x")
    ax_phi.set_ylabel("y")
    fig.colorbar(im_phi, ax=ax_phi)

    # Time series plot for averaged density and potential.
    ax_time.set_title("Time Series: Avg. Density & Avg. Potential (per cycle)")
    ax_time.set_xlabel("Time step")
    ax_time.set_ylabel("Value")
    time_line, = ax_time.plot([], [], 'b-', label="Avg Density")
    pot_line, = ax_time.plot([], [], 'r-', label="Avg Potential")
    ax_time.legend()

    # Particle positions plot.
    ax_part.set_title("Test Particle Positions (x-y)")
    ax_part.set_xlabel("x")
    ax_part.set_ylabel("y")
    colors = plt.cm.rainbow(np.linspace(0, 1, num_particles))
    particle_positions = np.random.uniform(-L/2, L/2, (num_particles, 3))
    particle_scatter = ax_part.scatter(particle_positions[:, 0], particle_positions[:, 1], c=colors)
    ax_part.set_xlim(-L/2, L/2)
    ax_part.set_ylim(-L/2, L/2)

    # Initialize simulation state.
    rho = np.zeros((N, N, N))
    particle_velocities = np.zeros((num_particles, 3))
    cycle_time = []
    cycle_avg_rho = []
    cycle_avg_phi = []
    global_step = 0
    current_cycle = 0

    def rk2_update(p, v, dt, acc_func):
        a1 = acc_func(p)
        p_temp = p + v * dt/2 + 0.5 * a1 * (dt/2)**2
        v_temp = v + a1 * dt/2
        a2 = acc_func(p_temp)
        v_new = v + dt * 0.5 * (a1 + a2)
        p_new = p + dt * v_new
        return p_new, v_new

    def update(frame):
        nonlocal global_step, current_cycle, rho, particle_positions, particle_velocities, target_shape

        if global_step > 0 and global_step % steps_per_cycle == 0:
            current_cycle += 1
            cycle_time.clear()
            cycle_avg_rho.clear()
            cycle_avg_phi.clear()
            rho[:] = 0.0
            particle_positions[:] = np.random.uniform(-L/2, L/2, (num_particles, 3))
            particle_velocities[:] = 0

        # (a) Add discrete collapse events.
        num_events = np.random.poisson(lam=collapse_rate)
        for _ in range(num_events):
            event_pos = np.random.uniform(-L/2, L/2, size=3)
            rho += collapse_amplitude * gaussian_3d(X, Y, Z,
                                                    event_pos[0],
                                                    event_pos[1],
                                                    event_pos[2],
                                                    collapse_sigma)
        # (b) Add continuous collapse noise.
        noise = continuous_noise_amplitude * np.random.randn(N, N, N)
        noise = gaussian_filter(noise, sigma=collapse_sigma/(L/N))
        rho += noise

        # (c) Apply density decay.
        rho *= density_decay

        # (d) Solve Poisson.
        phi = solve_poisson(rho, dx=(L/N), G=G)

        # (e) Compute gradient.
        interp_grad = interpolate_gradient(phi, grid, dx=(L/N))

        # (f) Update particles using RK2.
        for i in range(num_particles):
            pos = particle_positions[i]
            vel = particle_velocities[i]
            acc_func = lambda p: -np.array([interp_grad[d](p.reshape(1, 3))[0] for d in range(3)])
            new_pos, new_vel = rk2_update(pos, vel, dt, acc_func)
            particle_positions[i] = apply_periodic_bc(new_pos, L)
            particle_velocities[i] = new_vel

        # (g) Record summary data.
        dens_slice = rho[:, :, N//2]
        phi_slice = phi[:, :, N//2]
        avg_rho = np.mean(dens_slice)
        avg_phi = np.mean(phi_slice)
        cycle_time.append(global_step % steps_per_cycle)
        cycle_avg_rho.append(avg_rho)
        cycle_avg_phi.append(avg_phi)

        # (h) Update plots.
        im_rho.set_data(dens_slice)
        im_phi.set_data(phi_slice)
        time_line.set_data(np.array(cycle_time), np.array(cycle_avg_rho))
        pot_line.set_data(np.array(cycle_time), np.array(cycle_avg_phi))
        ax_time.relim()
        ax_time.autoscale_view()
        particle_scatter.set_offsets(particle_positions[:, :2])
        fig.suptitle(f"Cycle {current_cycle+1}/{num_cycles} - Step {(global_step % steps_per_cycle)+1}/{steps_per_cycle}")
        global_step += 1

        # Capture frame.
        fig.canvas.draw()
        frame_image = np.frombuffer(fig.canvas.tostring_rgb(), dtype='uint8')
        frame_image = frame_image.reshape(fig.canvas.get_width_height()[::-1] + (3,))
        if target_shape is None:
            target_shape = frame_image.shape
        else:
            if frame_image.shape != target_shape:
                frame_image = frame_image[:target_shape[0], :target_shape[1], :]
        video_frames.append(frame_image)
        return im_rho, im_phi, time_line, pot_line, particle_scatter

    anim = FuncAnimation(fig, update, frames=total_steps, interval=50, blit=False, repeat=False)
    plt.tight_layout()

    if interactive_mode:
        plt.show()
    else:
        plt.close(fig)

    # Save video if desired.
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    video_filename = f"visualization_simulation_{timestamp}.mp4"
    try:
        imageio.mimsave(video_filename, video_frames, fps=20)
        print(f"[INFO] Visualization video saved as {video_filename}")
    except Exception as e:
        print(f"[ERROR] Could not save visualization video: {e}")

###############################################################################
# 4. MAIN
###############################################################################

def main():
    # First, run the parameter sweep.
    print("[INFO] Starting parameter sweep...")
    param_sweep()
    print("[INFO] Parameter sweep completed.")

    # Ask user if they want to generate visualizations.
    generate_visuals = input("Do you want to run the visualization simulation? (y/n): ").strip().lower() == 'y'
    if generate_visuals:
        print("[INFO] Starting visualization simulation...")
        run_visualization_simulation()
        print("[INFO] Visualization simulation completed.")
    else:
        print("[INFO] Visualization simulation skipped.")

if __name__ == '__main__':
    main()

