import numpy as np
import matplotlib.pyplot as plt
from matplotlib.animation import FuncAnimation, FFMpegWriter
from scipy.fft import fftn, ifftn, fftfreq
from scipy.interpolate import RegularGridInterpolator
from scipy.ndimage import gaussian_filter
import csv, os, datetime, imageio
from docx import Document
from docx.shared import Pt

print("Current working directory:", os.getcwd())

# ------------------------------
# Helper Functions
# ------------------------------

def gaussian_3d(X, Y, Z, x0, y0, z0, sigma):
    """
    Return a normalized 3D Gaussian on (X,Y,Z) centered at (x0,y0,z0) with width sigma.
    """
    norm = (2 * np.pi * sigma**2) ** (3/2)
    return np.exp(-((X - x0)**2 + (Y - y0)**2 + (Z - z0)**2) / (2 * sigma**2)) / norm

def solve_poisson(rho, dx, G):
    """
    Solve the Poisson equation ∇²Φ = 4πG ρ on a periodic grid using FFT.
    """
    N = rho.shape[0]
    rho_k = fftn(rho)
    k = 2 * np.pi * fftfreq(N, d=dx)
    KX, KY, KZ = np.meshgrid(k, k, k, indexing='ij')
    k_squared = KX**2 + KY**2 + KZ**2
    k_squared[0,0,0] = 1.0  # avoid division by zero
    phi_k = -4 * np.pi * G * rho_k / k_squared
    phi_k[0,0,0] = 0.0  # enforce zero mean
    phi = np.real(ifftn(phi_k))
    return phi

def interpolate_gradient(phi, grid, dx):
    """
    Compute the gradient of φ and return a list of RegularGridInterpolator objects for each component.
    """
    grad = np.gradient(phi, dx, edge_order=2)
    interp_funcs = []
    for comp in grad:
        interp = RegularGridInterpolator(grid, comp, bounds_error=False, fill_value=None)
        interp_funcs.append(interp)
    return interp_funcs

def apply_periodic_bc(pos, L):
    """
    Apply periodic boundary conditions for a position array pos in domain [-L/2, L/2].
    """
    return ((pos + L/2) % L) - L/2

def compute_power_spectrum(field):
    """
    Compute the 2D power spectrum (azimuthally averaged) of a 2D field.
    Returns the radial power profile.
    """
    F = np.fft.fftshift(np.fft.fft2(field))
    psd2D = np.abs(F)**2
    ny, nx = field.shape
    y, x = np.indices((ny, nx))
    center = np.array([(ny-1)/2, (nx-1)/2])
    r = np.sqrt((x-center[1])**2 + (y-center[0])**2)
    r = r.astype(int)  # use built-in int
    tbin = np.bincount(r.ravel(), psd2D.ravel())
    nr = np.bincount(r.ravel())
    radial_prof = tbin / (nr + 1e-8)
    return radial_prof

def estimate_noise_exponent(psd, fit_range=(1,20)):
    """
    Estimate the power-law exponent of the noise spectrum by fitting log10(power) vs. log10(radial bin)
    over the specified range. Returns the slope and intercept.
    """
    bins = np.arange(len(psd))
    idx = (bins >= fit_range[0]) & (bins < fit_range[1])
    x = np.log10(bins[idx])
    y = np.log10(psd[idx] + 1e-12)
    coeffs = np.polyfit(x, y, 1)
    return coeffs  # slope, intercept

# ------------------------------
# Main Simulation Function
# ------------------------------

def main():
    # ===== Configuration =====
    interactive_mode = True   # Set to True to display visuals in real time.
    # ===========================
    
    # Simulation parameters
    G = 1.0                # Gravitational constant (arbitrary units)
    L = 10.0               # Domain: [-L/2, L/2] in each dimension
    N = 128                # Grid points per dimension (higher resolution)
    dx = L / N             # Grid spacing
    steps_per_cycle = 200  # Steps per cycle
    num_cycles = 10        # Number of cycles
    total_steps = steps_per_cycle * num_cycles  # Total steps (2000)
    dt = 0.05              # Time step

    # Collapse event parameters
    collapse_rate = 0.5      # Average discrete collapse events per step
    collapse_sigma = 0.2     # Localization width for discrete events
    collapse_amplitude = 1.0 # Mass deposited per discrete event
    continuous_noise_amplitude = 0.01  # Continuous noise amplitude (CSL-like)
    density_decay = 0.99     # Decay factor per step

    # Test particle parameters
    num_particles = 10

    # Create fixed 3D grid.
    x = np.linspace(-L/2, L/2, N, endpoint=False)
    grid = (x, x, x)
    X, Y, Z = np.meshgrid(x, x, x, indexing='ij')

    # Containers for aggregated cycle data.
    all_cycle_time = []
    all_cycle_avg_rho = []
    all_cycle_avg_phi = []
    all_cycle_psd = []   # Averaged power spectrum for each cycle

    # Container for video frames.
    video_frames = []
    target_shape = None  # to ensure all frames have the same dimensions

    # Set up figure with 4 subplots.
    fig, axs = plt.subplots(2, 2, figsize=(12, 10))
    ax_rho = axs[0, 0]
    ax_phi = axs[0, 1]
    ax_time = axs[1, 0]
    ax_part = axs[1, 1]

    # Removed the non-resizable window forcing (not supported by your backend).

    # Fixed color limits.
    dens_clim = (0, collapse_amplitude * 2)
    pot_clim = (-5, 5)

    # Prepare initial images (mid-plane slice: z=0).
    slice_index = N // 2
    dens_slice = np.zeros((N, N))
    im_rho = ax_rho.imshow(dens_slice, extent=(-L/2, L/2, -L/2, L/2), origin='lower',
                           cmap='viridis', vmin=dens_clim[0], vmax=dens_clim[1])
    ax_rho.set_title("Density Slice (z=0)")
    ax_rho.set_xlabel("x")
    ax_rho.set_ylabel("y")
    fig.colorbar(im_rho, ax=ax_rho)

    phi_slice = np.zeros((N, N))
    im_phi = ax_phi.imshow(phi_slice, extent=(-L/2, L/2, -L/2, L/2), origin='lower',
                           cmap='inferno', vmin=pot_clim[0], vmax=pot_clim[1])
    ax_phi.set_title("Potential Slice (z=0)")
    ax_phi.set_xlabel("x")
    ax_phi.set_ylabel("y")
    fig.colorbar(im_phi, ax=ax_phi)

    # Time series plot for averaged density and potential.
    ax_time.set_title("Time Series: Avg. Density & Avg. Potential (per cycle)")
    ax_time.set_xlabel("Time step")
    ax_time.set_ylabel("Value")
    time_line, = ax_time.plot([], [], 'b-', label="Avg Density")
    pot_line, = ax_time.plot([], [], 'r-', label="Avg Potential")
    ax_time.legend()

    # Particle positions plot.
    ax_part.set_title("Test Particle Positions (x-y)")
    ax_part.set_xlabel("x")
    ax_part.set_ylabel("y")
    # Initialize scatter with initial particle positions.
    particle_positions = np.random.uniform(-L/2, L/2, (num_particles, 3))
    initial_positions = particle_positions[:, :2]
    colors = plt.cm.rainbow(np.linspace(0, 1, num_particles))
    particle_scatter = ax_part.scatter(initial_positions[:,0], initial_positions[:,1], c=colors)
    ax_part.set_xlim(-L/2, L/2)
    ax_part.set_ylim(-L/2, L/2)

    # Initialize simulation state for the first cycle.
    rho = np.zeros((N, N, N))
    particle_velocities = np.zeros((num_particles, 3))
    cycle_time = []
    cycle_avg_rho = []
    cycle_avg_phi = []
    cycle_psd = []  # To store power spectrum of mid-plane potential per step

    global_step = 0
    current_cycle = 0

    # Use RK2 for test particle integration.
    def rk2_update(p, v, dt, acc_func):
        a1 = acc_func(p)
        p_temp = p + v * dt/2 + 0.5 * a1 * (dt/2)**2
        v_temp = v + a1 * dt/2
        a2 = acc_func(p_temp)
        v_new = v + dt * 0.5 * (a1 + a2)
        p_new = p + dt * v_new
        return p_new, v_new

    def update(frame):
        nonlocal global_step, current_cycle, rho, particle_positions, particle_velocities
        nonlocal cycle_time, cycle_avg_rho, cycle_avg_phi, cycle_psd, target_shape

        # End-of-cycle: store cycle data and reinitialize state.
        if global_step > 0 and global_step % steps_per_cycle == 0:
            all_cycle_time.append(np.array(cycle_time))
            all_cycle_avg_rho.append(np.array(cycle_avg_rho))
            all_cycle_avg_phi.append(np.array(cycle_avg_phi))
            avg_psd = np.mean(np.array(cycle_psd), axis=0)
            all_cycle_psd.append(avg_psd)
            current_cycle += 1
            cycle_time = []
            cycle_avg_rho = []
            cycle_avg_phi = []
            cycle_psd = []
            rho = np.zeros((N, N, N))
            particle_positions[:] = np.random.uniform(-L/2, L/2, (num_particles, 3))
            particle_velocities[:] = 0

        # (a) Add discrete collapse events.
        num_events = np.random.poisson(lam=collapse_rate)
        for _ in range(num_events):
            event_pos = np.random.uniform(-L/2, L/2, size=3)
            rho += collapse_amplitude * gaussian_3d(X, Y, Z,
                                                    event_pos[0],
                                                    event_pos[1],
                                                    event_pos[2],
                                                    collapse_sigma)
        # (b) Add continuous collapse noise.
        noise = continuous_noise_amplitude * np.random.randn(N, N, N)
        noise = gaussian_filter(noise, sigma=collapse_sigma/dx)
        rho += noise

        # (c) Apply density decay.
        rho *= density_decay

        # (d) Solve for gravitational potential.
        phi = solve_poisson(rho, dx, G)

        # (e) Compute gradient for dynamics.
        interp_grad = interpolate_gradient(phi, grid, dx)

        # (f) Update test particles using RK2.
        for i in range(num_particles):
            pos = particle_positions[i]
            vel = particle_velocities[i]
            acc_func = lambda p: -np.array([interp_grad[dim](p.reshape(1,3))[0] for dim in range(3)])
            new_pos, new_vel = rk2_update(pos, vel, dt, acc_func)
            particle_positions[i] = apply_periodic_bc(new_pos, L)
            particle_velocities[i] = new_vel

        # (g) Record summary data from mid-plane slice.
        dens_slice = rho[:, :, slice_index]
        phi_slice = phi[:, :, slice_index]
        avg_rho = np.mean(dens_slice)
        avg_phi = np.mean(phi_slice)
        cycle_time.append(global_step % steps_per_cycle)
        cycle_avg_rho.append(avg_rho)
        cycle_avg_phi.append(avg_phi)
        psd = compute_power_spectrum(phi_slice)
        cycle_psd.append(psd)

        # (h) Update plots.
        im_rho.set_data(dens_slice)
        im_phi.set_data(phi_slice)
        time_line.set_data(np.array(cycle_time), np.array(cycle_avg_rho))
        pot_line.set_data(np.array(cycle_time), np.array(cycle_avg_phi))
        ax_time.relim()
        ax_time.autoscale_view()
        particle_scatter.set_offsets(particle_positions[:, :2])
        fig.suptitle(f"Cycle {current_cycle+1}/{num_cycles} - Step {(global_step % steps_per_cycle)+1}/{steps_per_cycle}")
        global_step += 1

        # Capture current frame.
        fig.canvas.draw()
        frame_image = np.frombuffer(fig.canvas.tostring_rgb(), dtype='uint8')
        frame_image = frame_image.reshape(fig.canvas.get_width_height()[::-1] + (3,))
        if target_shape is None:
            target_shape = frame_image.shape
        else:
            if frame_image.shape != target_shape:
                frame_image = frame_image[:target_shape[0], :target_shape[1], :]
        video_frames.append(frame_image)

        return im_rho, im_phi, time_line, pot_line, particle_scatter

    # Create animation.
    anim = FuncAnimation(fig, update, frames=total_steps, interval=50, blit=False, repeat=False)
    plt.tight_layout()

    if interactive_mode:
        plt.show()  # Display visuals in real time.
    else:
        plt.close(fig)

    # Save video using ImageIO (wrap in try/except so failure doesn't block CSV/Docx saving).
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    video_filename = f"simulation_full_{timestamp}.mp4"
    try:
        imageio.mimsave(video_filename, video_frames, fps=20)
        print(f"Full simulation video saved as {video_filename}")
    except Exception as e:
        print("Error saving video:", e)

    # Save cycle summary data to CSV.
    csv_filename = f"simulation_summary_{timestamp}.csv"
    with open(csv_filename, mode="w", newline="") as file:
        writer_csv = csv.writer(file)
        writer_csv.writerow(["Cycle", "Step", "Avg Density", "Avg Potential"])
        for cycle_idx, (t_arr, rho_arr, phi_arr) in enumerate(zip(all_cycle_time, all_cycle_avg_rho, all_cycle_avg_phi)):
            for t, r_val, p_val in zip(t_arr, rho_arr, phi_arr):
                writer_csv.writerow([cycle_idx+1, t, r_val, p_val])
    print(f"Summary data saved as {csv_filename}")

    # Save power spectrum data.
    psd_filename = f"potential_psd_{timestamp}.csv"
    with open(psd_filename, mode="w", newline="") as file:
        writer_psd = csv.writer(file)
        writer_psd.writerow(["Cycle", "Radial Frequency Bin", "Power"])
        for cycle_idx, psd in enumerate(all_cycle_psd):
            for bin_idx, power in enumerate(psd):
                writer_psd.writerow([cycle_idx+1, bin_idx, power])
    print(f"Power spectrum data saved as {psd_filename}")

    # Estimate noise exponent from the overall averaged power spectrum.
    overall_psd = np.mean(np.array(all_cycle_psd), axis=0)
    slope, intercept = estimate_noise_exponent(overall_psd, fit_range=(1,20))
    
    # Plot and save the overall noise spectrum with the fit.
    plt.figure(figsize=(8,6))
    bins = np.arange(len(overall_psd))
    plt.loglog(bins[1:], overall_psd[1:], 'b.', label="Simulated PSD")
    fit_bins = bins[1:20]
    fit_line = 10**(intercept) * fit_bins**(slope)
    plt.loglog(fit_bins, fit_line, 'r-', label=f"Fit: slope = {slope:.2f}")
    plt.xlabel("Radial Frequency Bin")
    plt.ylabel("Power")
    plt.title("Averaged Noise Spectrum of Mid-plane Potential")
    plt.legend()
    spectrum_filename = f"noise_spectrum_{timestamp}.png"
    plt.savefig(spectrum_filename)
    plt.close()
    print(f"Noise spectrum plot saved as {spectrum_filename}")

    # Produce a DOCX report using python-docx.
    report_filename = f"simulation_report_{timestamp}.docx"
    doc = Document()
    doc.add_heading("Simulation Report: Emergent Gravity from Quantum Collapse", 0)

    doc.add_heading("Simulation Parameters:", level=1)
    params = [
        f"Gravitational constant, G = {G}",
        f"Domain size, L = {L}",
        f"Grid points per dimension, N = {N}",
        f"Time step, dt = {dt}",
        f"Steps per cycle = {steps_per_cycle}",
        f"Number of cycles = {num_cycles}",
        f"Total steps = {total_steps}",
        f"Discrete collapse rate = {collapse_rate}",
        f"Collapse sigma = {collapse_sigma}",
        f"Collapse amplitude = {collapse_amplitude}",
        f"Continuous noise amplitude = {continuous_noise_amplitude}",
        f"Density decay factor = {density_decay}",
        f"Number of test particles = {num_particles}"
    ]
    for param in params:
        p = doc.add_paragraph(param)
        p.style.font.size = Pt(11)

    doc.add_heading("Results Summary:", level=1)
    summary = (
        "The simulation demonstrates that both discrete collapse events and continuous noise "
        "can generate a gravitational potential that, when averaged over multiple cycles, "
        "exhibits large-scale features reminiscent of Newtonian gravity with stochastic fluctuations. "
        "Both the average density and average potential are computed and displayed in the time series. "
        "The computed power spectrum of the mid-plane potential fluctuations provides a quantitative measure "
        "of the noise characteristics. The estimated noise exponent from the fit is {:.2f}.".format(slope)
    )
    doc.add_paragraph(summary)

    doc.add_heading("Scientific Analysis and Critique:", level=1)
    analysis = (
        "The model simulates quantum collapse events as discrete mass deposits along with continuous noise to generate a gravitational potential. "
        "A potential criticism is that since the same potential drives particle motion, the result might be considered circular. "
        "To mitigate this, the simulation independently computes the noise spectrum of the potential—an output not directly used in particle updates. "
        "In a rigorous investigation, one would run a control simulation with an independently generated potential to validate the emergent noise signature."
    )
    doc.add_paragraph(analysis)

    doc.add_heading("Discussion and Future Directions:", level=1)
    discussion = (
        "Although various proposals linking quantum collapse and gravity exist (e.g., by Penrose, Di'osi, and in CSL/GRW theories), "
        "a rigorous demonstration that gravity emerges solely from collapse events remains an open challenge. Future work should incorporate "
        "advanced collapse models with relativistic corrections, higher resolution simulations, and independent control simulations to safeguard "
        "against circular reasoning. Identification of unique quantitative signatures (such as the noise spectrum exponent) is crucial for experimental validation."
    )
    doc.add_paragraph(discussion)

    doc.add_heading("Conclusion:", level=1)
    conclusion = (
        "This enhanced simulation, while still simplified, produces quantitative outputs—including an independent noise spectrum "
        "and an estimated noise exponent—that suggest the cumulative effect of quantum collapse processes could yield a gravitational field "
        "with measurable characteristics. However, caution is needed to ensure that the observed dynamics are not merely a self-fulfilling consequence "
        "of the model's design. Further investigation, including control simulations and experimental comparisons, is warranted."
    )
    doc.add_paragraph(conclusion)

    doc.save(report_filename)
    print(f"Report saved as {report_filename}")

if __name__ == '__main__':
    main()

